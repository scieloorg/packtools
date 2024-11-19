from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

from packtools.sps.formats.pdf.pipeline import xml as xml_pipe
from packtools.sps.formats.pdf.utils import docx_utils, xml_utils
from packtools.sps.formats.pdf import enum as pdf_enum


def pipeline_docx(xml_tree, data):
    """
    Create a DOCX document from an XML tree and additional data.

    Args:
        xml_tree: The XML tree containing the article data.
        data: Additional data for the DOCX generation.

    Returns:
        A DOCX Document object.
    """
    docx = docx_utils.init_docx(data)

    # First page header
    journal_title = xml_pipe.extract_journal_title(xml_tree)
    docx_journal_title_pipe(docx, journal_title)

    doi = xml_pipe.extract_doi(xml_tree)
    docx_doi_pipe(docx, doi)

    # First page content
    article_type = xml_pipe.extract_article_type(xml_tree)
    category = xml_pipe.extract_category(xml_tree)
    docx_article_type_and_category_pipe(docx, category, article_type)

    article_title = xml_pipe.extract_article_title(xml_tree)
    docx_article_title_pipe(docx, article_title)

    contrib_data = xml_pipe.extract_contrib_data(xml_tree)
    docx_authors_pipe(docx, contrib_data['authors_names'])
    docx_affiliation_pipe(docx, contrib_data['affiliations'])
    docx_corresponding_pipe(docx, contrib_data['corresponding_author'])

    article_main_language = xml_pipe.extract_article_main_language(xml_tree)
    abstract_data = xml_pipe.extract_abstract_data(xml_tree)
    docx_abstract_pipe(docx, abstract_data['title'], abstract_data['content'])

    keywords_data = xml_pipe.extract_keywords_data(xml_tree, article_main_language)
    docx_keyworks_pipe(docx, keywords_data['title'], keywords_data['keywords'])

    trans_abstract_data = xml_pipe.extract_trans_abstract_data(xml_tree)
    for ta in trans_abstract_data:
        docx_abstract_pipe(docx, ta['title'], ta['content'])
        ta_keywords_data = xml_pipe.extract_keywords_data(xml_tree, ta['lang'])
        docx_keyworks_pipe(docx, ta_keywords_data['title'], ta_keywords_data['keywords'])

    # Next pages header
    docx_second_header_pipe(docx, journal_title, article_title)

    # First page footer
    footer_data = xml_pipe.extract_footer_data(xml_tree)
    cite_as_part_one = xml_pipe.extract_cite_as_part_one(xml_tree) 
    docx_cite_as_pipe(docx, cite_as_part_one, journal_title, footer_data)
    docx_page_vol_issue_year_pipe(docx, footer_data)

    # Next Pages footer
    docx_second_footer_pipe(docx, footer_data)
    
    # Main content
    body_data = xml_pipe.extract_body_data(xml_tree)
    docx_body_pipe(docx, body_data)
    
    # Acknowledgments
    acknow_data = xml_pipe.extract_acknowledgment_data(xml_tree)
    docx_acknowledgments_pipe(docx, acknow_data['title'], acknow_data['paragraphs'])

    # References
    references_data = xml_pipe.extract_references_data(xml_tree)
    references = map(xml_utils.get_text_from_mixed_citation_node, references_data['references'])
    docx_references_pipe(docx, references_data['title'], references)

    # Supplementary Material
    supplementary_data = xml_pipe.extract_supplementary_data(xml_tree)
    if supplementary_data['elements']:
        docx_supplementary_material_pipe(docx, footer_data, supplementary_data)

    # Setting up sections
    docx_utils.docx_setup_sections(docx)
    
    return docx

def docx_journal_title_pipe(docx, journal_title_text, style_name='SCL Journal Title Char'):
    """
    Adds the journal title text to the first page header of the DOCX document, with each word on a new line.
    
    Args:
        docx (python-docx.Document): The DOCX document object.
        journal_title_text (str): The text of the journal title to be added.
        style_name (str, optional): The name of the style to apply to the journal title text. Defaults to 'SCL Journal Title Char'.
    
    Returns:
        python-docx.Paragraph: The paragraph object containing the journal title text.
    """
    first_page_header = docx_utils.get_first_page_header(docx)
    first_page_header
    para = docx_utils.get_first_paragraph(first_page_header)

    left_run = para.add_run(journal_title_text.replace(' ', '\n'))
    left_run.style = docx.styles[style_name]

    return para

def docx_doi_pipe(docx, doi_code, paragraph=None, style_name='SCL Header Paragraph Char'):
    """
    Adds the DOI (Digital Object Identifier) code to the first page header of the DOCX document, with the DOI URL formatted as a tab-indented string.
    
    Args:
        docx (python-docx.Document): The DOCX document object.
        doi_code (str): The DOI code to be added.
        paragraph (python-docx.Paragraph, optional): The paragraph object to add the DOI URL to. If not provided, the first paragraph in the first page header will be used.
        style_name (str, optional): The name of the style to apply to the DOI URL. Defaults to 'SCL Header Paragraph Char'.
    
    Returns:
        None
    """
    doi_url = f"http://dx.doi.org/{doi_code}"

    if paragraph:
        para = paragraph
    else:
        first_page_header = docx_utils.get_first_page_header(docx)
        para = docx_utils.get_first_paragraph(first_page_header)

    r = para.add_run(f'\t{doi_url}')
    r.style = docx.styles[style_name]

def docx_article_type_and_category_pipe(docx, category, article_type='Original Article', style_name='SCL Article Category'):
    """
    Adds the article category and type to the first page header of the DOCX document.
    Args:
        docx (python-docx.Document): The DOCX document object.
        category (str):
    
    Returns:
        None
    """
    article_category_text = ' | '.join([category, article_type])
    article_category_el = docx.add_paragraph(article_category_text)
    article_category_el.style = docx.styles[style_name]
