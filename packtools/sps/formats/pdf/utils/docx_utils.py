from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from packtools.sps.formats.pdf import enum as pdf_enum


def init_docx(data):
    """
    Initialize a DOCX document with optional base layout.

    Args:
        data: Additional data for the DOCX generation.

    Returns:
        A DOCX Document object.
    """
    source_layout = data.get('base_layout')
    if not source_layout:
        return Document()

    source_docx = Document(source_layout)
    docx = Document()
    copy_styles(source_docx, docx)

    return docx

def copy_styles(source, target):
    """
    Copy styles from a source DOCX document to a target DOCX document.

    Args:
        source: The source DOCX document.
        target: The target DOCX document.

    Returns:
        The target DOCX document with copied styles.
    """
    for style in source.styles:
        if style.type in [
            pdf_enum.WD_STYLE_TYPE.PARAGRAPH, 
            pdf_enum.WD_STYLE_TYPE.CHARACTER
        ]:
            if style.name in target.styles:
                target.styles[style.name].delete()
            
            new_style = target.styles.add_style(style.name, style.type)
            new_style.base_style = style.base_style
            new_style.hidden = style.hidden
            new_style.priority = style.priority
            new_style.quick_style = style.quick_style
            new_style.unhide_when_used = style.unhide_when_used

            if style.font.name:
                new_style.font.name = style.font.name
            
            if style.font.size:
                new_style.font.size = style.font.size
            
            new_style.font.bold = style.font.bold
            new_style.font.italic = style.font.italic
            new_style.font.underline = style.font.underline
            new_style.font.strike = style.font.strike
            new_style.font.all_caps = style.font.all_caps
            
            if style.font.color.rgb:
                new_style.font.color.rgb = style.font.color.rgb
            
            if style.font.color.theme_color:
                new_style.font.color.theme_color = style.font.color.theme_color
            
            new_style.font.highlight_color = style.font.highlight_color
            new_style.font.superscript = style.font.superscript
            new_style.font.subscript = style.font.subscript

            if style.type == pdf_enum.WD_STYLE_TYPE.PARAGRAPH:
                try:
                    new_style.paragraph_format.alignment = style.paragraph_format.alignment
                except ValueError:
                    new_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
                new_style.paragraph_format.right_indent = style.paragraph_format.right_indent
                new_style.paragraph_format.first_line_indent = style.paragraph_format.first_line_indent
                new_style.paragraph_format.keep_together = style.paragraph_format.keep_together
                new_style.paragraph_format.keep_with_next = style.paragraph_format.keep_with_next
                new_style.paragraph_format.page_break_before = style.paragraph_format.page_break_before
                new_style.paragraph_format.widow_control = style.paragraph_format.widow_control
                new_style.paragraph_format.space_before = style.paragraph_format.space_before
                new_style.paragraph_format.space_after = style.paragraph_format.space_after
                new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
                new_style.paragraph_format.line_spacing_rule = style.paragraph_format.line_spacing_rule

    return target

def docx_setup_sections(docx, page_attributes=pdf_enum.PAGE_ATTRIBUTES):
    """
    Sets up the sections of a DOCX document with the specified page attributes.
    
    Args:
        docx (python-docx.Document): The DOCX document object.
        page_attributes (dict, optional): A dictionary of page attributes to apply to the sections. Defaults to pdf_enum.PAGE_ATTRIBUTES.
    
    Returns:
        None
    """
    for ind, sec in enumerate(docx.sections):
        if ind == 0:
            sec.different_first_page_header_footer = page_attributes.get('different_first_page_header_footer')
        elif ind == 1:
            sec.header.is_linked_to_previous = False

        sec.top_margin = page_attributes.get('top_margin')
        sec.left_margin = page_attributes.get('left_margin')
        sec.right_margin = page_attributes.get('right_margin')
        sec.bottom_margin = page_attributes.get('bottom_margin')
        sec.header_distance = page_attributes.get('header_distance')
        sec.footer_distance = page_attributes.get('footer_distance')
        sec.gutter = page_attributes.get('gutter')
        sec.orientation = page_attributes.get('orientation')
        sec.page_height = page_attributes.get('page_height')
        sec.page_width = page_attributes.get('page_width')
        sec.start_type = page_attributes.get('start_type')

def get_first_section(docx):
    """
    Gets the first section of the given docx document.

    Args:
        docx (python-docx.Document): The docx document to get the first section from.

    Returns:
        python-docx.Section: The first section of the docx document.
    """
    try:
        return docx.sections[0]
    except IndexError:
        docx.add_section()
        return docx.sections[0]

def get_or_create_second_section(docx):
    """
    Get or create the second section of a DOCX document.

    If the document has only one section, a new section is added.
    Returns the second section of the document.

    Parameters:
    docx (Document): A python-docx Document object.

    Returns:
    Section: The second section of the DOCX document.
    """
    while len(docx.sections) < 2:
        docx.add_section()
    return docx.sections[1]

def get_first_page_header(docx):
    """
    Retrieves the first page header from the first section of a DOCX document.

    Args:
        docx: A DOCX document object.

    Returns:
        The first page header of the first section of the DOCX document.
    """
    section = get_first_section(docx)
    return section.first_page_header

def get_first_page_footer(docx):
    """
    Retrieves the first page footer from the first section of a DOCX document.

    Args:
        docx: A DOCX document object.

    Returns:
        The first page footer of the first section of the DOCX document.
    """
    section = get_first_section(docx)
    return section.first_page_footer

def get_second_header(docx):
    """
    Retrieves the header of the second section in a DOCX document.

    Args:
        docx (Document): A python-docx Document object representing the DOCX file.

    Returns:
        Header: The header of the second section in the DOCX document.
    """
    section = get_or_create_second_section(docx)
    return section.header

def get_second_footer(docx):
    """
    Retrieves the footer of the second section in a DOCX document.

    Args:
        docx (docx.Document): The DOCX document object.

    Returns:
        docx.section.SectionFooter: The footer of the second section in the DOCX document.
    """
    section = get_or_create_second_section(docx)
    return section.footer

def get_first_paragraph(element):
    """
    Retrieves the first paragraph from a given document element.

    Args:
        element: A document element that contains paragraphs.

    Returns:
        The first paragraph of the given element.
    """
    try:
        return element.paragraphs[0]
    except KeyError:
        element.add_paragraph()
        return element.paragraphs[0]

def add_table(docx, table_data, header_style_name='SCL Table Heading'):
    """
    Adds a table to a DOCX document with the specified data and formatting.

    Args:
        docx (docx.Document): The DOCX document to which the table will be added.
        table_data (dict): A dictionary containing the table data with the following keys:
            - 'label' (str): The label for the table.
            - 'title' (str): The title of the table.
            - 'headers' (list of list of str): A list containing the header rows of the table.
            - 'rows' (list of list of str): A list containing the data rows of the table.

    Returns:
        None
    """
    add_paragraph_with_formatting(docx, 
        f"{table_data['label']}. {table_data['title']}", 
        style_name=header_style_name,
    )

    headers = table_data['headers']
    rows = table_data['rows']

    num_cols = len(headers[0]) if headers else len(rows[0])
    table = docx.add_table(rows=1 + len(rows), cols=num_cols)

    if headers:
        header_row = table.rows[0]
        for i, header in enumerate(headers[0]):
            cell = header_row.cells[i]
            try:
                cell.text = header
            except TypeError:
                cell.text = ''
            style_cell(cell, bold=True, font_size=7, align='center')

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            try:
                cell = row.cells[j]
                cell.text = cell_data if cell_data else ''
                style_cell(cell, font_size=7, align='right')
            except IndexError:
                ...

    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading'

def style_cell(cell, bold=False, font_size=7, font_color=None, align='center', bg_color=None):
    """
    Styles a table cell in a DOCX document.

    Args:
        cell (docx.table._Cell): The cell to style.
        bold (bool, optional): If True, sets the text to bold. Defaults to False.
        font_size (int, optional): The font size of the text. Defaults to 7.
        font_color (docx.shared.RGBColor, optional): The font color as an RGBColor object. Defaults to None.
        align (str, optional): The alignment of the text. Can be 'center', 'left', or 'right'. Defaults to 'center'.
        bg_color (str, optional): The background color of the cell in hex format (e.g., 'FFFFFF' for white). Defaults to None.

    Returns:
        None
    """
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    font = run.font
    font.bold = bold
    font.size = Pt(font_size)
    
    if font_color:
        font.color.rgb = font_color
    
    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if bg_color:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), bg_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

def add_paragraph_with_formatting(docx, text, style_name='SCL Paragraph', element=None):
    """
    Adds a paragraph with specified formatting to a docx document.

    Args:
        docx (docx.Document): The docx document to which the paragraph will be added.
        text (str): The text content of the paragraph.
        style_name (str, optional): The name of the style to apply to the paragraph. Defaults to 'SCL Paragraph'.
        element (docx.oxml.CT_Element, optional): An optional element to which the paragraph will be added. 
            If None, the paragraph will be added to the docx document. Defaults to None.

    Returns:
        docx.text.paragraph.Paragraph: The newly added paragraph with the specified formatting.
    """
    if element is not None:
        para = element.add_paragraph(text)
    else:
        para = docx.add_paragraph(text)
    
    para.style = docx.styles[style_name]

    return para

def add_authors_names_paragraph_with_formatting_sup(docx, authors_names, paragraph_style_name, character_style_name, sup_mark="[^]"):
    """
    Adds a paragraph to a docx document with formatted author names and superscript markers.

    Args:
        docx (docx.Document): The docx document to which the paragraph will be added.
        authors_names (list of str): A list of author names, each potentially containing a superscript marker.
        paragraph_style_name (str): The style name to be applied to the paragraph.
        character_style_name (str): The style name to be applied to the text runs within the paragraph.
        sup_mark (str, optional): The marker used to split normal text and superscript text in author names. Defaults to "[^]".

    Returns:
        docx.text.paragraph.Paragraph: The paragraph object that was added to the document.
    """
    para = docx.add_paragraph()
    para.style = docx.styles[paragraph_style_name]

    for ind, a in enumerate(authors_names):
        try:
            normal_text, sup_text = a.split(sup_mark)
        except ValueError:
            normal_text = a
            sup_text = ''
                
        normal_run = para.add_run(normal_text)
        normal_run.style = docx.styles[character_style_name]
        
        separator_text = ''

        if ind <= len(authors_names) - 3:
            separator_text = ', '

        if ind == len(authors_names) - 2:
            separator_text = ' and '

        sup_run = para.add_run(sup_text)
        sup_run.style = docx.styles[character_style_name]
        sup_run.font.superscript = True

        sep_run = para.add_run(separator_text)
        sep_run.style = docx.styles[character_style_name]

    return para

def add_text_paragraph_with_formatting_sup(docx, text, paragraph_style_name, character_style_name, sup_mark="[^]"):
    """
    Adds a paragraph with formatted text to a docx document, including superscript formatting.
    
    Args:
        docx (docx.Document): The docx document to which the paragraph will be added.
        text (str): The text to be added to the paragraph. The text should contain a superscript marker.
        paragraph_style_name (str): The name of the paragraph style to be applied.
        character_style_name (str): The name of the character style to be applied.
        sup_mark (str, optional): The marker indicating the superscript text. Defaults to "[^]".
    
    Raises:
        ValueError: If the text does not contain the superscript marker.
    """
    para = docx.add_paragraph()
    para.style = docx.styles[paragraph_style_name]

    try:
        sup_text, normal_text = text.split(sup_mark)
    except ValueError:
        normal_text = text
        sup_text = ''
            
    sup_run = para.add_run(sup_text)
    sup_run.style = docx.styles[character_style_name]
    sup_run.font.superscript = True

    normal_run = para.add_run(normal_text)
    normal_run.style = docx.styles[character_style_name]

def add_run_with_style(element, text, style):
    run = element.add_run(text)
    run.style = style

def add_heading_with_formatting(docx, text, style_name, level):
    """
    Adds a heading to a DOCX document with specified formatting.

    Args:
        docx (docx.Document): The DOCX document to which the heading will be added.
        text (str): The text content of the heading.
        style_name (str): The name of the style to apply to the heading.
        level (int): The level of the heading (e.g., 0 for title, 1 for main heading, etc.).

    Returns:
        docx.text.paragraph.Paragraph: The newly added heading with the specified formatting
    """
    heading = docx.add_heading(text, level=level)
    heading.style = docx.styles[style_name]

    return heading

def level_to_style(level):
    """
    Convert a heading level to a corresponding style name.

    Args:
        level (int): The heading level to convert.

    Returns:
        str: The corresponding style name. Returns 'SCL Section Title' for level 2,
             'SCL Subsection Title' for level 3, and 'SCL Paragraph' for all other levels.
    """
    if level == 2:
        return 'SCL Section Title'
    elif level == 3:
        return 'SCL Subsection Title'
    else:
        return 'SCL Paragraph'

def setup_section_columns(section, num_columns, column_spacing):
    """
    Configures the number of columns and the spacing between them for a given section in a DOCX document.

    Args:
        section (docx.section.Section): The section of the DOCX document to configure.
        num_columns (int): The number of columns to set in the section.
        column_spacing (int): The spacing between columns in twips (twentieths of a point).

    Returns:
        None
    """
    sect_pr = section._sectPr
    cols = sect_pr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sect_pr.append(cols)
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), str(column_spacing))
