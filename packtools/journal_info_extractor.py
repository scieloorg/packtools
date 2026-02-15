#!/usr/bin/env python3
# coding: utf-8
"""
Journal Information Extractor

Extrai informações de documentos DOCX de periódicos científicos e gera planilhas XLSX.
Processa documentos multilíngues (português, inglês, espanhol) normalizando seções para inglês.

Gera duas planilhas:
1. sections.xlsx - Todas as seções extraídas dos documentos
2. editorial_board.xlsx - Dados estruturados do corpo editorial

Usage:
    journal-extractor <input_dir> [options]
    journal-extractor --help
"""

import os
import sys
import re
import argparse
import logging
from pathlib import Path
from datetime import datetime
from collections import OrderedDict

from docx import Document
from openpyxl import Workbook

LOGGER = logging.getLogger(__name__)

# ============================================================================
# MAPEAMENTO DE SEÇÕES PT/ES -> EN
# ============================================================================

SECTION_MAPPING = {
    # Seções principais
    "SOBRE O PERIÓDICO": "ABOUT THE JOURNAL",
    "POLÍTICA EDITORIAL": "EDITORIAL POLICY",
    "CORPO EDITORIAL": "EDITORIAL BOARD",
    "INSTRUÇÕES PARA OS AUTORES": "INSTRUCTIONS FOR AUTHORS",
    "INSTRUÇÕES PARA AUTORES": "INSTRUCTIONS FOR AUTHORS",

    # Subseções de ABOUT THE JOURNAL
    "Breve Histórico": "Brief History",
    "Acesso Aberto": "Open Access",
    "Conformidade com a Ciência Aberta": "Open Science Compliance",
    "Ética na Publicação": "Publication Ethics",
    "Foco e Escopo": "Focus and Scope",
    "Preservação Digital": "Digital Preservation",
    "Fontes de Indexação": "Indexing Sources",
    "Ficha Bibliográfica": "Bibliographic Record",
    "Sites e Redes Sociais": "Websites and Social Media",
    "Websites e Mídias Sociais": "Websites and Social Media",

    # Subseções de EDITORIAL POLICY
    "Preprints": "Preprints",
    "Processo de Avaliação por Pares": "Peer Review Process",
    "Processo de avaliação por pares": "Peer Review Process",
    "Dados Abertos": "Open Data",
    "Dados abertos": "Open Data",
    "Taxas de Artigo": "Article Fees",
    "Cobrança de Taxas": "Article Fees",
    "Ética, Más Condutas, Erratas e Retratações": "Ethics, Misconduct, Errata, and Retractions",
    "Política de Ética e Más condutas, Errata e Retratação": "Ethics, Misconduct, Errata, and Retractions",
    "Política de Conflito de Interesse": "Conflict of Interest Policy",
    "Política sobre Conflito de Interesses": "Conflict of Interest Policy",
    "Uso de Software de Verificação de Similaridade": "Use of Similarity-Checking Software",
    "Adoção de softwares de verificação de similaridade": "Use of Similarity-Checking Software",
    "Uso de Ferramentas de Inteligência Artificial": "Use of Artificial Intelligence Tools",
    "Adoção de softwares uso de recursos de Inteligência Artificial": "Use of Artificial Intelligence Tools",
    "Uso por Autores": "Use by Authors",
    "Uso por autores": "Use by Authors",
    "Responsabilidade e Transparência": "Accountability and Transparency",
    "Responsabilidade e transparência": "Accountability and Transparency",
    "Uso por Revisores e Editores": "Use by Reviewers and Editors",
    "Uso por pareceristas e editores": "Use by Reviewers and Editors",
    "Processos de Avaliação e Decisões Editoriais": "Evaluation Processes and Editorial Decisions",
    "Processos de avaliação e decisões editoriais": "Evaluation Processes and Editorial Decisions",
    "Atualizações": "Updates",
    "Questões de Sexo e Gênero": "Sex and Gender Issues",
    "Comitê de Ética": "Ethics Committee",
    "Direitos Autorais": "Copyright",
    "Propriedade Intelectual e Termos de Uso": "Intellectual Property and Terms of Use",
    "Propriedade Intelectual e Termos de uso": "Intellectual Property and Terms of Use",
    "Responsabilidade do Site": "Website Responsibility",
    "Responsabilidade do site": "Website Responsibility",
    "Responsabilidade do Autor": "Author Responsibility",
    "Responsabilidade do autor": "Author Responsibility",
    "Patrocinadores e Agências de Fomento": "Sponsors and Funding Agencies",

    # Subseções de EDITORIAL BOARD
    "Editor-Chefe": "Editor-in-Chief",
    "Editor-in-Chief": "Editor-in-Chief",  # Inglês -> Inglês
    "Editores Executivos": "Executive Editors",
    "Executive Editors": "Executive Editors",  # Inglês -> Inglês
    "Editores Associados: Física Teórica, Física Computacional e Temas de Fronteira": "Associate Editors: Theoretical Physics, Computational Physics and Frontier Topics",
    "Editores Associados: Física Teórica, Computacional e Temas de Fronteira": "Associate Editors: Theoretical Physics, Computational Physics and Frontier Topics",
    "Associate Editors: Theoretical Physics, Computational Physics and Frontier Topics": "Associate Editors: Theoretical Physics, Computational Physics and Frontier Topics",
    # Inglês -> Inglês
    "Editores Associados: Física Experimental": "Associate Editors: Experimental Physics",
    "Associate Editors: Experimental Physics": "Associate Editors: Experimental Physics",  # Inglês -> Inglês
    "Editores Associados: Pesquisa em Ensino de Física": "Associate Editors: Physics Education Research",
    "Editoras Associadas: Pesquisa em Ensino de Física": "Associate Editors: Physics Education Research",
    "Associate Editors: Physics Education Research": "Associate Editors: Physics Education Research",
    # Inglês -> Inglês
    "Editores Associados: Epistemologia e História da Física e da Astronomia": "Associate Editors: Epistemology and History of Physics and Astronomy",
    "Editores Associados: Epistemologia e História da Física e Astronomia": "Associate Editors: Epistemology and History of Physics and Astronomy",
    "Associate Editors: Epistemology and History of Physics and Astronomy": "Associate Editors: Epistemology and History of Physics and Astronomy",
    # Inglês -> Inglês
    "Editores Honorários": "Honorary Editors",
    "Honorary Editors": "Honorary Editors",  # Inglês -> Inglês

    # Subseções de INSTRUCTIONS FOR AUTHORS
    "Tipos de Submissões Aceitas": "Types of Accepted Submissions",
    "Tipos de documentos aceitos": "Types of Accepted Submissions",
    "Contribuições dos Autores": "Author Contributions",
    "Contribuição dos Autores": "Author Contributions",
    "Formato de Submissão de Artigos": "Article Submission Format",
    "Formato de Envio dos Artigos": "Article Submission Format",
    "Ativos Digitais": "Digital Assets",
    "Citações e Referências": "Citations and References",
    "Declaração de Financiamento": "Funding Declaration",
    "Informações Adicionais": "Additional Information",
    "Informações de Contato": "Contact Information",
    "Contato": "Contact Information",

    # Seções especiais de Open Data
    'Se disponível no próprio artigo': 'Open Data',
    'Se disponível em repositório': 'Open Data',
    'Se disponível anonimizado': 'Open Data',
    'Se disponível mediante solicitação ao autor correspondente': 'Open Data',
    'Se disponível mediante solicitação a organização': 'Open Data',
    'If available in the article itself': 'Open Data',
    'If available in a repository': 'Open Data',
    'If available anonymized': 'Open Data',
    'If available upon request from the corresponding author': 'Open Data',
    'If available upon request from an organization': 'Open Data',

    # Seções especiais de Citations
    "Journal Article": "Citations and References",
    "Artigo de periódico": "Citations and References",
    "Periódico": "Citations and References",
    "Book": "Citations and References",
    "Livro": "Citations and References",
    "Book Chapter": "Citations and References",
    "Capítulo de Livro": "Citations and References",
    "Capítulo de livro": "Citations and References",
    "Proceedings": "Citations and References",
    "Anais": "Citations and References",
    "Thesis": "Citations and References",
    "Tese": "Citations and References",
    "Teses": "Citations and References",
    "Preprint": "Citations and References",

    # Seções de tipos de artigos
    "Artigos Gerais": "General Articles",
    "Produtos e Materiais Didáticos para o Ensino de Física": "Products and Didactic Materials for Physics Teaching",
    "Pesquisa em Ensino de Física": "Research in Physics Education",
    "História da Física e Ciências Afins": "History of Physics and Related Sciences",
}


# ============================================================================
# EXCEPTION CLASSES
# ============================================================================

class JournalExtractorError(Exception):
    """Base exception for journal extractor errors."""
    pass


class InvalidFileError(JournalExtractorError):
    """Exception raised when file cannot be processed."""
    pass


class NoFilesFoundError(JournalExtractorError):
    """Exception raised when no valid files are found."""
    pass


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def extract_metadata_from_filename(filename):
    """
    Extrai metadados do nome do arquivo.

    Padrão esperado: YYYYMMDD_ACRONIMO_.*_IDIOMA_ok.docx

    Args:
        filename (str): Nome do arquivo

    Returns:
        tuple: (data_iso, acronimo, idioma_code) ou (None, None, None) se não encontrar

    Examples:
        >>> extract_metadata_from_filename("20251010_RBEF_Total_Página_Informativa_inglês_ok.docx")
        ('2025-10-10', 'RBEF', 'en')
    """
    pattern = r'(\d{8})_([A-Z]+)_.*_(inglês|português|espanhol|english|portuguese|spanish)_ok\.docx'
    match = re.search(pattern, filename, re.IGNORECASE)

    if match:
        data = match.group(1)
        acronimo = match.group(2)
        idioma_raw = match.group(3).lower()

        # Converter para ISO 639-1
        idioma_map = {
            'inglês': 'en',
            'english': 'en',
            'português': 'pt',
            'portuguese': 'pt',
            'espanhol': 'es',
            'spanish': 'es'
        }
        idioma = idioma_map.get(idioma_raw, idioma_raw)

        # Formatar data para ISO (YYYY-MM-DD)
        data_iso = f"{data[:4]}-{data[4:6]}-{data[6:]}"

        return data_iso, acronimo, idioma

    return None, None, None


def extract_hyperlink_url(paragraph):
    """
    Extrai URLs de hyperlinks de um parágrafo.

    Args:
        paragraph: Objeto Paragraph do python-docx

    Returns:
        list: Lista de URLs encontradas
    """
    urls = []
    if paragraph._element.xml:
        for hyperlink in paragraph._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
            r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if r_id:
                try:
                    url = paragraph.part.rels[r_id].target_ref
                    urls.append(url)
                except Exception as e:
                    LOGGER.debug(f"Error extracting hyperlink: {e}")
    return urls


def get_paragraph_text_with_urls(paragraph):
    """
    Obtém o texto do parágrafo incluindo URLs de hyperlinks (sem tags HTML).

    Args:
        paragraph: Objeto Paragraph do python-docx

    Returns:
        str: Texto do parágrafo com URLs
    """
    text = paragraph.text.strip()
    urls = extract_hyperlink_url(paragraph)

    if urls:
        clean_urls = []
        for url in urls:
            # Remover mailto: prefix de emails
            if url.startswith('mailto:'):
                clean_urls.append(url.replace('mailto:', ''))
            else:
                clean_urls.append(url)
        text = text + " " + " ".join(clean_urls)

    return text


def is_section_header(paragraph):
    """
    Verifica se um parágrafo é um cabeçalho de seção.

    Considera tanto parágrafos em negrito quanto estilos Heading.
    Distingue cabeçalhos de conteúdo formatado (campo: valor).

    Args:
        paragraph: Objeto Paragraph do python-docx

    Returns:
        bool: True se for cabeçalho de seção
    """
    if not paragraph.text.strip():
        return False

    text = paragraph.text.strip()

    # Verificar se é um estilo de cabeçalho (Heading 1, Heading 2, etc.)
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith('Heading'):
        return True

    # Verificar se tem negrito
    has_bold = any(run.bold for run in paragraph.runs if run.text.strip())

    if not has_bold:
        return False

    # Se tem negrito mas é formato "campo: valor", é conteúdo, não cabeçalho
    # EXCETO se for um role de Associate/Associate Editors ou Editores Associados
    if ':' in text:
        # Verificar se é um role de Associate/Associate Editors
        if (text.startswith("Associate Editors:") or
                text.startswith("Editores Associados:") or
                text.startswith("Editoras Associadas:")):
            return True  # É um role válido

        # Verificar se tem conteúdo após os dois pontos
        parts = text.split(':', 1)
        if len(parts) == 2 and parts[1].strip():
            # É um campo com valor, não é cabeçalho de seção
            return False

    return True


def normalize_section_name(section_name, is_main_section=False):
    """
    Normaliza o nome da seção para inglês.

    Se já estiver em inglês, retorna como está.
    Se estiver em português/espanhol, traduz usando o mapeamento.

    Args:
        section_name (str): Nome da seção original
        is_main_section (bool): Se é seção principal

    Returns:
        str: Nome normalizado em inglês
    """
    # Correspondência exata
    if section_name in SECTION_MAPPING:
        return SECTION_MAPPING[section_name]

    # Correspondência parcial - procurar a chave mais longa que corresponde
    best_match = None
    best_match_len = 0

    for pt, en in SECTION_MAPPING.items():
        # Se a seção começa com a chave do mapeamento
        if section_name.startswith(pt) and len(pt) > best_match_len:
            best_match = en
            best_match_len = len(pt)

    if best_match:
        return best_match

    # Se não encontrar e for seção principal em português, tentar identificar
    if is_main_section:
        section_upper = section_name.upper()
        for pt, en in SECTION_MAPPING.items():
            if pt.upper() == section_upper:
                return en

    # Caso contrário, retornar como está (sem modificações)
    # NÃO cortar roles como "Associate Editors: Theoretical Physics..."
    return section_name


# ============================================================================
# SECTION EXTRACTOR CLASS
# ============================================================================

class SectionExtractor:
    """
    Extrai seções de documentos DOCX multilíngues.

    Attributes:
        docx_path (Path): Caminho do arquivo DOCX
        sections (OrderedDict): Seções extraídas (nome_en: conteúdo)
    """

    # Seções principais (apenas 4)
    MAIN_SECTION_KEYWORDS = {
        "ABOUT THE JOURNAL", "SOBRE O PERIÓDICO",
        "EDITORIAL POLICY", "POLÍTICA EDITORIAL",
        "EDITORIAL BOARD", "CORPO EDITORIAL",
        "INSTRUCTIONS FOR AUTHORS", "INSTRUÇÕES PARA OS AUTORES", "INSTRUÇÕES PARA AUTORES"
    }

    # Seções a serem ignoradas (processadas separadamente)
    SKIP_SECTIONS = {"EDITORIAL BOARD", "CORPO EDITORIAL"}

    def __init__(self, docx_path):
        """
        Inicializa o extrator de seções.

        Args:
            docx_path (str or Path): Caminho do arquivo DOCX

        Raises:
            InvalidFileError: Se o arquivo não puder ser lido
        """
        self.docx_path = Path(docx_path)
        self.sections = OrderedDict()

        if not self.docx_path.exists():
            raise InvalidFileError(f"File not found: {docx_path}")

    def extract(self):
        """
        Extrai todas as seções do documento.

        EXCLUI a seção Editorial Board (será processada separadamente).

        Returns:
            OrderedDict: Seções extraídas {nome_en: conteúdo}

        Raises:
            InvalidFileError: Se houver erro ao processar o documento
        """
        try:
            doc = Document(self.docx_path)
        except Exception as e:
            raise InvalidFileError(f"Error reading {self.docx_path}: {e}")

        current_section = None
        current_content = []
        skip_until_subsection = False
        skip_editorial_board = False

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            if is_section_header(para):
                # Verificar se é Editorial Board (pular completamente)
                if text in self.SKIP_SECTIONS or text.upper() in self.SKIP_SECTIONS:
                    skip_editorial_board = True
                    # Salvar seção anterior antes de pular
                    if current_section and current_content:
                        self._save_section(current_section, current_content)
                        current_content = []
                    current_section = None
                    continue

                # Verificar se saímos do Editorial Board
                if skip_editorial_board and text in self.MAIN_SECTION_KEYWORDS:
                    skip_editorial_board = False

                # Se estamos no Editorial Board, pular tudo
                if skip_editorial_board:
                    continue

                # Normalizar nome da seção
                normalized_section = normalize_section_name(text)

                # Verificar se é seção principal
                if text in self.MAIN_SECTION_KEYWORDS or text.upper() in self.MAIN_SECTION_KEYWORDS:
                    # Salvar seção anterior
                    if current_section and current_content:
                        self._save_section(current_section, current_content)
                        current_content = []

                    current_section = None
                    skip_until_subsection = True
                else:
                    # É subseção
                    if current_section and current_section != normalized_section and current_content:
                        self._save_section(current_section, current_content)
                        current_content = []
                    elif current_section == normalized_section:
                        # Mesma seção, não limpar conteúdo (vai concatenar)
                        pass
                    else:
                        # Nova seção, limpar conteúdo se havia algo
                        if current_content:
                            current_content = []

                    current_section = normalized_section
                    skip_until_subsection = False
            else:
                # Se estamos no Editorial Board, pular
                if skip_editorial_board:
                    continue

                # Conteúdo da seção
                if current_section and not skip_until_subsection:
                    text_with_urls = get_paragraph_text_with_urls(para)
                    if text_with_urls:
                        current_content.append(text_with_urls)

        # Salvar última seção
        if current_section and current_content:
            self._save_section(current_section, current_content)

        LOGGER.info(f"Extracted {len(self.sections)} sections from {self.docx_path.name}")
        return self.sections

    def _save_section(self, section_name, content):
        """
        Salva ou concatena conteúdo de uma seção.

        Args:
            section_name (str): Nome normalizado da seção
            content (list): Lista de strings com o conteúdo
        """
        content_str = "\n".join(content).strip()
        if section_name in self.sections:
            # Concatenar com conteúdo existente
            self.sections[section_name] += "\n" + content_str
        else:
            self.sections[section_name] = content_str

    def get_journal_info(self):
        """
        Extrai title_journal e issn_scielo da seção Bibliographic Record.

        Returns:
            tuple: (title_journal, issn_scielo)
        """
        title_journal = ""
        issn_scielo = ""

        if "Bibliographic Record" in self.sections:
            content = self.sections["Bibliographic Record"]
            lines = content.split("\n")

            for line in lines:
                line = line.strip()
                # Extrair Journal Title / Título do periódico
                if line.startswith("Journal Title:") or line.startswith("Título do periódico:"):
                    title_journal = line.split(":", 1)[1].strip()
                # Extrair ISSN
                elif line.startswith("ISSN:"):
                    issn_scielo = line.split(":", 1)[1].strip()

        return title_journal, issn_scielo


# ============================================================================
# EDITORIAL BOARD EXTRACTOR CLASS
# ============================================================================

class EditorialBoardExtractor:
    """
    Extrai informações do corpo editorial de documentos DOCX.

    Attributes:
        docx_path (Path): Caminho do arquivo DOCX
        members (list): Lista de dicionários com dados dos membros
    """

    def __init__(self, docx_path):
        """
        Inicializa o extrator de corpo editorial.

        Args:
            docx_path (str or Path): Caminho do arquivo DOCX

        Raises:
            InvalidFileError: Se o arquivo não puder ser lido
        """
        self.docx_path = Path(docx_path)
        self.members = []

        if not self.docx_path.exists():
            raise InvalidFileError(f"File not found: {docx_path}")

    def extract(self):
        """
        Extrai todos os membros do corpo editorial usando lógica robusta.

        Returns:
            list: Lista de dicionários com dados dos membros

        Raises:
            InvalidFileError: Se houver erro ao processar o documento
        """
        try:
            doc = Document(self.docx_path)
        except Exception as e:
            raise InvalidFileError(f"Error reading {self.docx_path}: {e}")

        current_role = None
        in_editorial_board = False

        # Keywords que indicam seção de Editorial Board
        editorial_board_keywords = ["EDITORIAL BOARD", "CORPO EDITORIAL"]

        # Keywords que indicam fim da seção
        end_keywords = ["INSTRUCTIONS FOR AUTHORS", "INSTRUÇÕES PARA OS AUTORES", "INSTRUÇÕES PARA AUTORES"]

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Verificar se entramos na seção Editorial Board
            if text in editorial_board_keywords:
                in_editorial_board = True
                LOGGER.debug(f"Entered Editorial Board section")
                continue

            # Verificar se saímos da seção Editorial Board
            if text in end_keywords:
                in_editorial_board = False
                LOGGER.debug(f"Left Editorial Board section")
                break

            if not in_editorial_board:
                continue

            # Se é um título de cargo (negrito)
            if is_section_header(para):
                current_role = normalize_section_name(text)
                LOGGER.debug(f"Found role: {current_role}")
                continue

            # Se é uma linha com informações de membro (contém vírgula e ponto)
            if current_role and "," in text and "." in text:
                member = self._parse_member(text, current_role, para)
                if member:
                    self.members.append(member)
                    LOGGER.debug(f"Extracted member: {member['name']} {member['surname']}")

        LOGGER.info(f"Extracted {len(self.members)} editorial board members from {self.docx_path.name}")
        return self.members

    def _parse_member(self, text, role, paragraph):
        """
        Faz parsing robusto de um membro do corpo editorial.

        Args:
            text (str): Texto do parágrafo
            role (str): Função/papel do membro
            paragraph: Objeto Paragraph do python-docx

        Returns:
            dict: Dados do membro
        """
        member = {
            'role': role,
            'name': '',
            'surname': '',
            'institution': '',
            'city': '',
            'state': '',
            'state_code': '',
            'country': 'Brasil',  # Default
            'country_code': 'BR',  # Default
            'lattes': '',
            'orcid': '',
            'email': ''
        }

        # Extrair URLs de hyperlinks
        urls = extract_hyperlink_url(paragraph)

        # Parse do texto
        # Formato: Nome Sobrenome, Instituição, Cidade, Estado/Código, País.
        # Exemplo: Carlos Eduardo Aguiar, Instituto de Física da UFRJ, Rio de Janeiro, RJ, Brasil.

        # Separar por vírgulas
        parts = [p.strip() for p in text.split(",")]

        # Parte 1: Nome completo
        if len(parts) >= 1:
            full_name = parts[0].strip()
            # Separar nome e sobrenome
            name_parts = full_name.split()
            if len(name_parts) >= 2:
                member['name'] = name_parts[0]
                member['surname'] = " ".join(name_parts[1:])
            else:
                member['name'] = full_name

        # Identificar onde termina a localização e começam os extras (Lattes, ORCID, e-mail)
        location_end_idx = len(parts)
        for i, part in enumerate(parts):
            if "Lattes" in part or "ORCID" in part or "e-mail" in part:
                location_end_idx = i
                # Se esta parte contém "Brasil." ou "Brazil." antes de Lattes/ORCID/e-mail,
                # ainda é parte da localização
                if ("Brasil." in part or "Brazil." in part) and i > 0:
                    part_clean = part.strip()
                    if part_clean.startswith("Brasil") or part_clean.startswith("Brazil"):
                        location_end_idx = i + 1
                break

        # Processar partes de localização DE TRÁS PARA FRENTE
        # Padrão: Nome, [Instituição (pode ter várias partes)], Cidade, Estado/Código, País
        location_parts = parts[1:location_end_idx]  # Remover nome

        if len(location_parts) == 0:
            return member

        # Processar de trás para frente
        # Última parte: País
        if len(location_parts) >= 1:
            country_part = location_parts[-1].strip()
            # Remover tudo após o primeiro ponto (pode ter ". Lattes:" etc)
            if "." in country_part:
                country_part = country_part.split(".")[0].strip()
            member['country'] = country_part
            if "Brasil" in country_part or "Brazil" in country_part:
                member['country_code'] = 'BR'

        # Verificar quantas partes temos para determinar o formato
        if len(location_parts) == 4:
            # Formato sem cidade: Instituição, Estado, Código, País
            # Penúltima parte: Código do Estado (2 letras)
            if len(location_parts[-2].strip()) == 2 and location_parts[-2].strip().isupper():
                member['state_code'] = location_parts[-2].strip()
                member['state'] = location_parts[-3].strip()
                member['institution'] = location_parts[0].strip()
                member['city'] = ''
            else:
                # Penúltima pode ter barra (Estado/Código)
                state_part = location_parts[-2].strip()
                if "/" in state_part:
                    state, state_code = state_part.split("/", 1)
                    member['state'] = state.strip()
                    member['state_code'] = state_code.strip()
                else:
                    member['state'] = state_part

                member['city'] = location_parts[-3].strip()
                member['institution'] = location_parts[0].strip()

        elif len(location_parts) >= 5:
            # Formato com cidade: Instituição(s), Cidade, Estado/Código, País
            state_part = location_parts[-2].strip()
            if "/" in state_part:
                state, state_code = state_part.split("/", 1)
                member['state'] = state.strip()
                member['state_code'] = state_code.strip()
            else:
                member['state'] = state_part

            # Antepenúltima parte: Cidade
            member['city'] = location_parts[-3].strip()

            # Tudo que sobrou: Instituição (pode ter várias vírgulas)
            institution_parts = location_parts[:-3]
            member['institution'] = ", ".join(institution_parts).strip()

        elif len(location_parts) == 3:
            # Formato mínimo: Instituição, Estado/Código, País
            state_part = location_parts[-2].strip()
            if "/" in state_part:
                state, state_code = state_part.split("/", 1)
                member['state'] = state.strip()
                member['state_code'] = state_code.strip()
            else:
                member['state'] = state_part

            member['institution'] = location_parts[0].strip()

        elif len(location_parts) >= 1:
            # Só tem instituição e país
            member['institution'] = location_parts[0].strip() if len(location_parts) > 1 else ''

        # Extrair Lattes, ORCID, e-mail
        text_lower = text.lower()

        # Lattes
        if "lattes:" in text_lower or "lattes," in text_lower:
            lattes_found = False
            for url in urls:
                if "lattes.cnpq.br" in url:
                    member['lattes'] = url
                    lattes_found = True
                    break

            if not lattes_found and "lattes:" in text_lower:
                try:
                    lattes_part = text.split("Lattes:")[1].split(",")[0].strip()
                    if "http" in lattes_part:
                        url = lattes_part.split()[0]
                        member['lattes'] = url
                except Exception as e:
                    LOGGER.debug(f"Error extracting Lattes: {e}")

        # ORCID
        if "orcid:" in text_lower:
            orcid_found = False
            for url in urls:
                if "orcid.org" in url:
                    member['orcid'] = url
                    orcid_found = True
                    break

            if not orcid_found:
                try:
                    orcid_part = text.split("ORCID:")[1].split(",")[0].strip()
                    if "http" in orcid_part or "0000-" in orcid_part:
                        url = orcid_part.split()[0]
                        if not url.startswith("http"):
                            url = f"https://orcid.org/{url}"
                        member['orcid'] = url
                except Exception as e:
                    LOGGER.debug(f"Error extracting ORCID: {e}")

        # E-mail
        if "e-mail:" in text_lower:
            email_found = False
            for url in urls:
                if "mailto:" in url:
                    email = url.replace("mailto:", "")
                    member['email'] = email
                    email_found = True
                    break

            if not email_found:
                try:
                    email_part = text.split("e-mail:")[-1].strip()
                    email = email_part.replace("_", "").strip()
                    if "." in email:
                        email_clean = ""
                        for char in email:
                            if char == ".":
                                if "@" in email_clean:
                                    email_clean += char
                                else:
                                    break
                            else:
                                email_clean += char
                        if "@" in email_clean:
                            member['email'] = email_clean
                except Exception as e:
                    LOGGER.debug(f"Error extracting email: {e}")

        return member


# ============================================================================
# XLSX REPORT GENERATOR CLASS
# ============================================================================

class XLSXReportGenerator:
    """
    Gera relatórios em formato XLSX.

    Attributes:
        output_dir (Path): Diretório de saída
        timestamp (str): Timestamp para nomes de arquivo
    """

    def __init__(self, output_dir):
        """
        Inicializa o gerador de relatórios.

        Args:
            output_dir (str or Path): Diretório onde salvar os relatórios
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.timestamp = datetime.now().strftime("%Y%m%dT%H%M%S")

    def create_sections_report(self, file_data_list):
        """
        Cria planilha com seções extraídas.

        Args:
            file_data_list (list): Lista de dicts com dados dos arquivos processados

        Returns:
            Path: Caminho do arquivo gerado
        """
        # Coletar todas as seções (mantendo ordem do primeiro arquivo)
        all_sections = OrderedDict()

        for idx, file_info in enumerate(file_data_list):
            if idx == 0:
                all_sections = file_info['sections'].copy()
            else:
                for section_name in file_info['sections'].keys():
                    if section_name not in all_sections:
                        all_sections[section_name] = None

        # Criar planilha
        wb = Workbook()
        ws = wb.active
        ws.title = "Sections"

        # Linhas de metadados (linhas 1-3)
        ws.cell(row=1, column=1, value="Date")
        ws.cell(row=2, column=1, value="Acronym")
        ws.cell(row=3, column=1, value="Language")

        # Preencher metadados de cada arquivo
        for col_idx, file_info in enumerate(file_data_list, start=2):
            ws.cell(row=1, column=col_idx, value=file_info['date'])
            ws.cell(row=2, column=col_idx, value=file_info['acronym'])
            ws.cell(row=3, column=col_idx, value=file_info['language'])

        # Escrever seções e conteúdo (a partir da linha 4)
        row_idx = 4
        for section_name in all_sections.keys():
            # Nome da seção (coluna A)
            ws.cell(row=row_idx, column=1, value=section_name)

            # Conteúdo de cada arquivo
            for col_idx, file_info in enumerate(file_data_list, start=2):
                content = file_info['sections'].get(section_name, "")
                ws.cell(row=row_idx, column=col_idx, value=content)

            row_idx += 1

        # Ajustar largura das colunas
        ws.column_dimensions['A'].width = 50
        for col_idx in range(2, len(file_data_list) + 2):
            col_letter = chr(64 + col_idx)
            ws.column_dimensions[col_letter].width = 60

        # Salvar
        output_path = self.output_dir / f"{self.timestamp}-sections.xlsx"
        wb.save(output_path)

        LOGGER.info(f"Sections report saved: {output_path}")
        return output_path

    def create_editorial_board_report(self, all_members):
        """
        Cria planilha com dados do corpo editorial.

        Args:
            all_members (list): Lista de dicts com dados dos membros

        Returns:
            Path: Caminho do arquivo gerado
        """
        # Remover duplicatas
        unique_members = []
        seen = set()

        for member in all_members:
            key = (
                member.get('name', ''),
                member.get('surname', ''),
                member.get('role', ''),
                member.get('date', ''),
                member.get('acronym', '')
            )
            if key not in seen:
                unique_members.append(member)
                seen.add(key)

        # Criar planilha
        wb = Workbook()
        ws = wb.active
        ws.title = "Editorial Board"

        # Cabeçalhos - CAMPOS OBRIGATÓRIOS
        headers = [
            "title_journal",  # Obrigatório
            "issn_scielo",  # Obrigatório (print ou electronic)
            "affiliation",  # Obrigatório (instituição)
            "given_names",  # Obrigatório (nome)
            "last_name",  # Obrigatório (sobrenome)
            "country_code",  # Obrigatório
            "state_name",  # Obrigatório
            "city_name",  # Obrigatório
            "std_role",  # Obrigatório (papel padronizado)
            # Campos opcionais
            "orcid",
            "lattes",
            "email"
        ]

        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col_idx, value=header)

        # Escrever dados dos membros (a partir da linha 2)
        row_idx = 2
        for member in unique_members:
            ws.cell(row=row_idx, column=1, value=member.get('title_journal', ''))
            ws.cell(row=row_idx, column=2, value=member.get('issn_scielo', ''))
            ws.cell(row=row_idx, column=3, value=member.get('institution', ''))  # affiliation
            ws.cell(row=row_idx, column=4, value=member.get('name', ''))  # given_names
            ws.cell(row=row_idx, column=5, value=member.get('surname', ''))  # last_name
            ws.cell(row=row_idx, column=6, value=member.get('country_code', ''))
            ws.cell(row=row_idx, column=7, value=member.get('state', ''))  # state_name
            ws.cell(row=row_idx, column=8, value=member.get('city', ''))  # city_name
            ws.cell(row=row_idx, column=9, value=member.get('role', ''))  # std_role

            # ORCID - criar hiperlink se houver URL
            if member.get('orcid'):
                cell = ws.cell(row=row_idx, column=10)
                cell.value = member['orcid']
                cell.hyperlink = member['orcid']
                cell.style = 'Hyperlink'

            # Lattes - criar hiperlink se houver URL
            if member.get('lattes'):
                cell = ws.cell(row=row_idx, column=11)
                cell.value = member['lattes']
                cell.hyperlink = member['lattes']
                cell.style = 'Hyperlink'

            # Email - criar hiperlink se houver email
            if member.get('email'):
                cell = ws.cell(row=row_idx, column=12)
                cell.value = member['email']
                cell.hyperlink = f"mailto:{member['email']}"
                cell.style = 'Hyperlink'

            row_idx += 1

        # Ajustar largura das colunas
        ws.column_dimensions['A'].width = 50  # title_journal
        ws.column_dimensions['B'].width = 15  # issn_scielo
        ws.column_dimensions['C'].width = 50  # affiliation
        ws.column_dimensions['D'].width = 20  # given_names
        ws.column_dimensions['E'].width = 30  # last_name
        ws.column_dimensions['F'].width = 12  # country_code
        ws.column_dimensions['G'].width = 25  # state_name
        ws.column_dimensions['H'].width = 20  # city_name
        ws.column_dimensions['I'].width = 50  # std_role
        ws.column_dimensions['J'].width = 50  # orcid
        ws.column_dimensions['K'].width = 50  # lattes
        ws.column_dimensions['L'].width = 40  # email

        # Salvar
        output_path = self.output_dir / f"{self.timestamp}-editorial_board.xlsx"
        wb.save(output_path)

        LOGGER.info(f"Editorial Board report saved: {output_path}")
        return output_path


# ============================================================================
# DOCUMENT PROCESSOR CLASS
# ============================================================================

class DocumentProcessor:
    """
    Processa múltiplos documentos DOCX e coordena a extração de dados.

    Attributes:
        input_dir (Path): Diretório com os arquivos DOCX
        output_dir (Path): Diretório para salvar relatórios
        docx_files (list): Lista de arquivos DOCX encontrados
    """

    def __init__(self, input_dir, output_dir):
        """
        Inicializa o processador de documentos.

        Args:
            input_dir (str or Path): Diretório com arquivos DOCX
            output_dir (str or Path): Diretório para salvar relatórios

        Raises:
            NoFilesFoundError: Se nenhum arquivo válido for encontrado
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)

        if not self.input_dir.exists():
            raise NoFilesFoundError(f"Input directory not found: {input_dir}")

        # Encontrar arquivos DOCX
        self.docx_files = sorted(list(self.input_dir.glob("*_ok.docx")))

        if not self.docx_files:
            raise NoFilesFoundError(f"No *_ok.docx files found in {input_dir}")

        LOGGER.info(f"Found {len(self.docx_files)} DOCX files to process")

    def process(self):
        """
        Processa todos os documentos e gera relatórios.

        Returns:
            dict: Caminhos dos relatórios gerados
        """
        file_data_list = []
        all_editorial_members = []

        # Processar cada arquivo
        for idx, docx_file in enumerate(self.docx_files, 1):
            LOGGER.info(f"[{idx}/{len(self.docx_files)}] Processing: {docx_file.name}")

            try:
                # Extrair metadados do nome do arquivo
                date, acronym, language = extract_metadata_from_filename(docx_file.name)

                if not date:
                    LOGGER.warning(f"Could not extract metadata from filename: {docx_file.name}")
                    continue

                # Extrair seções
                section_extractor = SectionExtractor(docx_file)
                sections = section_extractor.extract()

                # Extrair informações do periódico
                title_journal, issn_scielo = section_extractor.get_journal_info()

                # Guardar dados do arquivo
                file_data_list.append({
                    'filename': docx_file.name,
                    'date': date,
                    'acronym': acronym,
                    'language': language,
                    'sections': sections,
                    'title_journal': title_journal,
                    'issn_scielo': issn_scielo
                })

                # Extrair corpo editorial
                try:
                    editorial_extractor = EditorialBoardExtractor(docx_file)
                    members = editorial_extractor.extract()

                    # Adicionar metadados e informações do periódico a cada membro
                    for member in members:
                        member['date'] = date
                        member['acronym'] = acronym
                        member['title_journal'] = title_journal
                        member['issn_scielo'] = issn_scielo
                        all_editorial_members.append(member)

                except Exception as e:
                    LOGGER.error(f"Error extracting editorial board from {docx_file.name}: {e}")

            except Exception as e:
                LOGGER.error(f"Error processing {docx_file.name}: {e}")
                continue

        if not file_data_list:
            raise NoFilesFoundError("No valid files were processed")

        # Gerar relatórios
        report_generator = XLSXReportGenerator(self.output_dir)

        sections_report = report_generator.create_sections_report(file_data_list)
        editorial_report = report_generator.create_editorial_board_report(all_editorial_members)

        return {
            'sections': sections_report,
            'editorial_board': editorial_report,
            'files_processed': len(file_data_list),
            'editorial_members': len(all_editorial_members)
        }


# ============================================================================
# MAIN FUNCTION
# ============================================================================

def main():
    """Entry point for journal-extractor CLI."""

    parser = argparse.ArgumentParser(
        description='Extract journal information from DOCX files and generate XLSX reports',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  journal-extractor /path/to/docx_files
  journal-extractor /path/to/docx_files --output /path/to/output
  journal-extractor /path/to/docx_files --loglevel DEBUG

Input files must follow naming pattern:
  YYYYMMDD_ACRONYM_.*_LANGUAGE_ok.docx
  Example: 20251010_RBEF_Total_Página_Informativa_inglês_ok.docx

Output files:
  YYYYMMDDTHHMMSS-sections.xlsx
  YYYYMMDDTHHMMSS-editorial_board.xlsx
        """
    )

    parser.add_argument(
        'input_dir',
        type=str,
        help='Directory containing DOCX files (*_ok.docx)'
    )

    parser.add_argument(
        '--output',
        '-o',
        type=str,
        default=None,
        help='Output directory for XLSX reports (default: current directory)'
    )

    parser.add_argument(
        '--loglevel',
        default='INFO',
        choices=['DEBUG', 'INFO', 'WARNING', 'ERROR', 'CRITICAL'],
        help='Set logging level (default: INFO)'
    )

    args = parser.parse_args()

    # Configurar logging
    logging.basicConfig(
        level=getattr(logging, args.loglevel.upper()),
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    # Definir diretório de saída
    output_dir = args.output if args.output else os.getcwd()

    try:
        # Processar documentos
        processor = DocumentProcessor(args.input_dir, output_dir)
        results = processor.process()

        # Exibir resultados
        print("\n" + "=" * 80)
        print("EXTRACTION COMPLETED SUCCESSFULLY")
        print("=" * 80)
        print(f"\nFiles processed: {results['files_processed']}")
        print(f"Editorial members extracted: {results['editorial_members']}")
        print(f"\nGenerated reports:")
        print(f"  - Sections: {results['sections']}")
        print(f"  - Editorial Board: {results['editorial_board']}")
        print("\n" + "=" * 80)

        return 0

    except (NoFilesFoundError, InvalidFileError, JournalExtractorError) as e:
        LOGGER.error(f"Error: {e}")
        return 1

    except Exception as e:
        LOGGER.exception(f"Unexpected error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())