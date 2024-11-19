import unittest

from docx import Document
from docx.shared import RGBColor

from packtools.sps.formats.pdf import enum as pdf_enum
from packtools.sps.formats.pdf.utils import docx_utils


class TestDocxUtils(unittest.TestCase):

    def setUp(self):
        self.doc = Document()
        self.doc.add_section()
        self.first_paragraph_text = 'This is the first paragraph.'
        self.doc.paragraphs[0].text = self.first_paragraph_text

    def test_get_first_section_is_not_none(self):
        first_section = docx_utils.get_first_section(self.doc)
        self.assertIsNotNone(first_section)

    def test_get_or_create_second_section_is_not_none(self):
        second_section = docx_utils.get_or_create_second_section(self.doc)
        self.assertIsNotNone(second_section)

    def test_get_first_page_header_is_not_none(self):
        header = docx_utils.get_first_page_header(self.doc)
        self.assertIsNotNone(header)

    def test_get_second_header_is_not_none(self):
        header = docx_utils.get_second_header(self.doc)
        self.assertIsNotNone(header)

    def test_get_first_page_footer_is_not_none(self):
        footer = docx_utils.get_first_page_footer(self.doc)
        self.assertIsNotNone(footer)

    def test_get_second_footer_is_not_none(self):
        footer = docx_utils.get_second_footer(self.doc)
        self.assertIsNotNone(footer)

    def test_get_first_paragraph(self):
        first_paragraph = docx_utils.get_first_paragraph(self.doc)
        self.assertEqual(self.first_paragraph_text, first_paragraph.text)

    def test_style_cell(self):
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "Test"
        docx_utils.style_cell(cell, bold=True, font_size=12, font_color=RGBColor(0xFF, 0x00, 0x00), align='center', bg_color='FFFF00')
        self.assertEqual(cell.paragraphs[0].runs[0].font.bold, True)
        self.assertEqual(cell.paragraphs[0].runs[0].font.size.pt, 12)
        self.assertEqual(cell.paragraphs[0].alignment, 1)  # Center alignment

    def test_add_table(self):
        my_new_table_paragraph_style_name = 'My new table paragraph style'
        self.doc.styles.add_style(my_new_table_paragraph_style_name, pdf_enum.WD_STYLE_TYPE.PARAGRAPH)

        table_data = {
            'label': 'Table 1',
            'title': 'Test Table',
            'headers': [['Header 1', 'Header 2']],
            'rows': [['Row 1 Col 1', 'Row 1 Col 2'], ['Row 2 Col 1', 'Row 2 Col 2']]
        }
        docx_utils.add_table(self.doc, table_data, my_new_table_paragraph_style_name)
        table = self.doc.tables[0]
        self.assertEqual(table.cell(0, 0).text, 'Header 1')
        self.assertEqual(table.cell(1, 0).text, 'Row 1 Col 1')

    def test_add_paragraph_with_formatting(self):
        my_new_paragraph_name = 'My new paragraph style'
        my_new_paragraph_style = self.doc.styles.add_style(my_new_paragraph_name, pdf_enum.WD_STYLE_TYPE.PARAGRAPH)
        para = docx_utils.add_paragraph_with_formatting(self.doc, "Test paragraph", style_name=my_new_paragraph_style.name)
        self.assertEqual(para.text, "Test paragraph")
        self.assertEqual(para.style.name, my_new_paragraph_name)

    def test_add_authors_names_paragraph_with_formatting_sup(self):
        # TODO
        ...

    def test_add_text_paragraph_with_formatting_sup(self):
        # TODO
        ...

    def test_add_run_with_style(self):
        new_style = self.doc.styles.add_style('My new character style', pdf_enum.WD_STYLE_TYPE.CHARACTER)

        para = self.doc.add_paragraph('This paragraph is about multimodal data and its style is `Default Paragraph Font`.')
        self.assertEqual(para.runs[0].text, "This paragraph is about multimodal data and its style is `Default Paragraph Font`.")
        self.assertEqual(para.runs[0].style.name, 'Default Paragraph Font')

        docx_utils.add_run_with_style(para, "This text has a very specific style! I am not normal.", new_style)
        self.assertEqual(para.runs[1].text, "This text has a very specific style! I am not normal.")
        self.assertEqual(para.runs[1].style.name, 'My new character style')

    def test_add_heading_with_formatting(self):
        heading = docx_utils.add_heading_with_formatting(self.doc, "Introduction to Data Mining", 'Heading 1', level=1)
        self.assertEqual(heading.text, "Introduction to Data Mining")
