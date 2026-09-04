import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from packtools.sps.formats.pdf.pipeline import docx as docx_pipe
from packtools.sps.formats.pdf.renderer import docx as docx_renderer
from packtools.sps.formats.pdf import enum as pdf_enum
from packtools.sps.utils import xml_utils

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

FIXTURES_DIR = Path(__file__).resolve().parents[4] / "fixtures" / "pdf"


def _docx_with_layout_styles():
    """A fresh Document carrying the named styles the footer/cite-as pipes rely on."""
    return docx_renderer.builder.init_docx({"base_layout": str(FIXTURES_DIR / "layout.docx")})


class TestPipelineDocx(unittest.TestCase):

    def _start_page_number(self, section):
        pg_num_type = section._sectPr.find(f'{WML_NS}pgNumType')
        return pg_num_type.get(f'{WML_NS}start')

    def _pipeline_docx(self, xml_filename):
        xml_tree = xml_utils.get_xml_tree(str(FIXTURES_DIR / xml_filename))
        data = {"base_layout": str(FIXTURES_DIR / "layout.docx"), "assets_dir": str(FIXTURES_DIR)}
        return docx_pipe.pipeline_docx(xml_tree, data)

    def test_start_page_number_uses_fpage_when_present(self):
        docx = self._pipeline_docx("a1.xml")
        self.assertEqual(self._start_page_number(docx.sections[0]), '271')

    def test_start_page_number_defaults_to_one_without_fpage(self):
        docx = self._pipeline_docx("a4.xml")
        self.assertEqual(self._start_page_number(docx.sections[0]), '1')


class TestFormatJournalTitleTwoLines(unittest.TestCase):
    """
    Regression tests for issue #1301: journal titles with more than 2 words
    used to get one word per line (unbounded), pushing the rest of the page-1
    header down. The first word stays on its own line and every remaining
    word is joined onto a single second line, capping the masthead at 2 lines
    regardless of how many words the title has.
    """

    def test_single_word_title_stays_on_one_line(self):
        self.assertEqual(docx_pipe._format_journal_title_two_lines('Biology'), 'Biology')

    def test_two_word_title_keeps_one_word_per_line(self):
        self.assertEqual(
            docx_pipe._format_journal_title_two_lines('Acta Amazonica'),
            'Acta\nAmazonica',
        )

    def test_four_word_title_is_capped_at_two_lines(self):
        self.assertEqual(
            docx_pipe._format_journal_title_two_lines('Brazilian Journal of Biology'),
            'Brazilian\nJournal of Biology',
        )

    def test_five_word_title_is_capped_at_two_lines(self):
        self.assertEqual(
            docx_pipe._format_journal_title_two_lines('Urbe. Revista Brasileira de Gestão Urbana'),
            'Urbe.\nRevista Brasileira de Gestão Urbana',
        )

    def test_empty_title_returns_empty_string(self):
        self.assertEqual(docx_pipe._format_journal_title_two_lines(''), '')


class TestJournalTitlePipe(unittest.TestCase):

    def setUp(self):
        self.docx = Document()
        self.docx.styles.add_style('SCL Journal Title Char', WD_STYLE_TYPE.CHARACTER)

    def test_two_word_title_keeps_one_word_per_line(self):
        para = docx_pipe.docx_journal_title_pipe(self.docx, 'Acta Amazonica')
        self.assertEqual(para.runs[0].text, 'Acta\nAmazonica')

    def test_multi_word_title_is_capped_at_two_lines(self):
        para = docx_pipe.docx_journal_title_pipe(self.docx, 'Brazilian Journal of Biology')
        self.assertEqual(para.runs[0].text, 'Brazilian\nJournal of Biology')

    def test_run_uses_the_given_style(self):
        para = docx_pipe.docx_journal_title_pipe(self.docx, 'Acta Amazonica')
        self.assertEqual(para.runs[0].style.name, 'SCL Journal Title Char')


class TestDocxDoiPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxArticleTypeAndCategoryPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxArticleTitlePipe(unittest.TestCase):
    # TODO
    ...


class TestDocxAuthorsPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxAffiliationPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxCorrespondingPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxAbstractPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxKeyworksPipe(unittest.TestCase):
    """
    Regression tests for issue #1322: the Keywords/Palavras-chave paragraph
    had no space-after of its own, so the section title immediately
    following it (space-before=0) collapsed onto it with almost no gap.
    """

    def setUp(self):
        self.docx = Document()
        self.docx.styles.add_style('SCL Paragraph Keywords', WD_STYLE_TYPE.PARAGRAPH)
        self.docx.styles.add_style('SCL Paragraph Keywords Header Char', WD_STYLE_TYPE.CHARACTER)
        self.docx.styles.add_style('SCL Paragraph Keywords Char', WD_STYLE_TYPE.CHARACTER)

    def test_paragraph_has_space_after(self):
        docx_pipe.docx_keyworks_pipe(self.docx, 'Keywords:', 'one, two, three')
        para = self.docx.paragraphs[-1]
        self.assertGreater(para.paragraph_format.space_after.pt, 0)

    def test_title_and_content_runs(self):
        docx_pipe.docx_keyworks_pipe(self.docx, 'Keywords:', 'one, two, three')
        para = self.docx.paragraphs[-1]
        self.assertEqual(para.runs[0].text, 'Keywords: ')
        self.assertEqual(para.runs[1].text, 'one, two, three')


class TestDocxCiteAsPipe(unittest.TestCase):

    def test_uses_fpage_lpage_range_when_present(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '10', 'issue': '2', 'year': '2023',
                       'fpage': 123, 'lpage': 130, 'location_label': '123-130'}
        docx_pipe.docx_cite_as_pipe(docx, 'Author AB. ', 'Journal Title', footer_data)

        footer = docx_renderer.section.get_first_page_footer(docx)
        para = docx_renderer.text.get_first_paragraph(footer)
        self.assertIn('10: 123-130.', para.text)

    def test_uses_elocation_id_when_fpage_is_absent(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '33', 'issue': '3', 'year': '2024',
                       'fpage': '', 'lpage': '', 'location_label': 'e282794'}
        docx_pipe.docx_cite_as_pipe(docx, 'Author AB. ', 'Journal Title', footer_data)

        footer = docx_renderer.section.get_first_page_footer(docx)
        para = docx_renderer.text.get_first_paragraph(footer)
        self.assertIn('33: e282794.', para.text)
        self.assertNotIn('-.', para.text)


class TestDocxSecondHeaderPipe(unittest.TestCase):

    def setUp(self):
        self.docx = Document()
        self.docx.styles.add_style('SCL Header Paragraph', WD_STYLE_TYPE.PARAGRAPH)
        self.docx.styles.add_style('SCL Header Paragraph Char', WD_STYLE_TYPE.CHARACTER)
        self.docx.styles.add_style('SCL Journal Title Char', WD_STYLE_TYPE.CHARACTER)

    def _second_header_paragraph(self):
        header = docx_pipe.docx_renderer.section.get_default_header(self.docx)
        return header.paragraphs[-1]

    def test_two_word_title_keeps_one_word_per_line(self):
        docx_pipe.docx_second_header_pipe(self.docx, 'Acta Amazonica', 'Some Article Title')
        para = self._second_header_paragraph()
        self.assertEqual(para.runs[0].text, 'Acta\nAmazonica')

    def test_multi_word_title_is_capped_at_two_lines(self):
        docx_pipe.docx_second_header_pipe(self.docx, 'Brazilian Journal of Biology', 'Some Article Title')
        para = self._second_header_paragraph()
        self.assertEqual(para.runs[0].text, 'Brazilian\nJournal of Biology')

    def test_article_title_is_appended_after_a_tab(self):
        docx_pipe.docx_second_header_pipe(self.docx, 'Acta Amazonica', 'Some Article Title')
        para = self._second_header_paragraph()
        self.assertEqual(para.runs[1].text, '\tSome Article Title')


class TestDocxSecondFooterPipe(unittest.TestCase):

    def test_uses_fpage_lpage_range_when_present(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '10', 'issue': '2', 'year': '2023',
                       'fpage': 123, 'lpage': 130, 'location_label': '123-130'}
        docx_pipe.docx_second_footer_pipe(docx, footer_data)

        footer = docx_renderer.section.get_second_footer(docx)
        para = footer.paragraphs[0]
        self.assertIn('VOL. 10 (2) 2023: 123-130', para.text)

    def test_uses_elocation_id_when_fpage_is_absent(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '33', 'issue': '3', 'year': '2024',
                       'fpage': '', 'lpage': '', 'location_label': 'e282794'}
        docx_pipe.docx_second_footer_pipe(docx, footer_data)

        footer = docx_renderer.section.get_second_footer(docx)
        para = footer.paragraphs[0]
        self.assertIn('VOL. 33 (3) 2024: e282794', para.text)
        self.assertNotIn(': -', para.text)


class TestDocxPageVolIssueYearPipe(unittest.TestCase):

    def test_uses_fpage_lpage_range_when_present(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '10', 'issue': '2', 'year': '2023',
                       'fpage': 123, 'lpage': 130, 'location_label': '123-130'}
        docx_pipe.docx_page_vol_issue_year_pipe(docx, footer_data)

        footer = docx_renderer.section.get_first_page_footer(docx)
        para = footer.paragraphs[-1]
        self.assertIn('VOL. 10 (2) 2023: 123-130', para.text)

    def test_uses_elocation_id_when_fpage_is_absent(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '33', 'issue': '3', 'year': '2024',
                       'fpage': '', 'lpage': '', 'location_label': 'e282794'}
        docx_pipe.docx_page_vol_issue_year_pipe(docx, footer_data)

        footer = docx_renderer.section.get_first_page_footer(docx)
        para = footer.paragraphs[-1]
        self.assertIn('VOL. 33 (3) 2024: e282794', para.text)
        self.assertNotIn(': -', para.text)


class TestDocxBodyPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxReferencesPipe(unittest.TestCase):
    """
    Regression: SCL Paragraph Reference uses the same font size as body
    text and has no line_spacing of its own, so a reference wrapping to
    multiple lines rendered with the same loose spacing as a body
    paragraph. Line spacing is now set explicitly on each reference
    paragraph.
    """

    def test_reference_has_single_line_spacing(self):
        docx = _docx_with_layout_styles()
        docx_pipe.docx_references_pipe(docx, references=['Author A. Title B. Journal C. 2020.'])
        para = docx.paragraphs[-1]
        self.assertEqual(para.paragraph_format.line_spacing, 1.0)


class TestDocxAcknowledgmentsPipe(unittest.TestCase):
    # TODO
    ...


class TestDocxSupplementaryMaterialPipe(unittest.TestCase):

    def test_footer_has_no_leading_pipe(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '53', 'issue': '4', 'year': '2023',
                       'fpage': 271, 'lpage': 280, 'location_label': '271-280'}
        docx_pipe.docx_supplementary_material_pipe(
            docx, footer_data, {'title': 'Supplementary Material', 'elements': []}
        )

        footer = docx.sections[-1].footer
        para = footer.paragraphs[-1]
        self.assertEqual(para.text, 'VOL. 53 (4) 2023: 271-280')

    def test_uses_elocation_id_when_fpage_is_absent(self):
        docx = _docx_with_layout_styles()
        footer_data = {'volume': '33', 'issue': '3', 'year': '2024',
                       'fpage': '', 'lpage': '', 'location_label': 'e282794'}
        docx_pipe.docx_supplementary_material_pipe(
            docx, footer_data, {'title': 'Supplementary Material', 'elements': []}
        )

        footer = docx.sections[-1].footer
        para = footer.paragraphs[-1]
        self.assertEqual(para.text, 'VOL. 33 (3) 2024: e282794')


class TestBodyColumnConfiguration(unittest.TestCase):
    """
    Regression tests: body column count (and the count restored after a
    full-width table/figure) must come from PAGE_ATTRIBUTES['default_column_count']
    instead of a hardcoded 2, so a 1-column body isn't forced back to 2
    columns after a full-width table/figure interlude.
    """

    def _cols_num(self, section):
        cols = section._sectPr.find(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols'
        )
        return cols.get(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'
        )

    def test_body_column_count_defaults_to_two(self):
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {}, clear=False):
            pdf_enum.PAGE_ATTRIBUTES.pop('default_column_count', None)
            self.assertEqual(docx_pipe._body_column_count(), 2)

    def test_body_column_count_reads_config(self):
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 1}):
            self.assertEqual(docx_pipe._body_column_count(), 1)

    def test_setup_body_section_uses_configured_column_count(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 1}):
            docx_pipe._setup_body_section(docx)
        self.assertEqual(self._cols_num(docx.sections[1]), '1')

    def test_setup_body_section_supports_three_columns(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 3}):
            docx_pipe._setup_body_section(docx)
        self.assertEqual(self._cols_num(docx.sections[1]), '3')

    def test_restore_body_column_section_restores_configured_column_count(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 1}):
            section = docx_pipe._restore_body_column_section(docx)
        self.assertEqual(self._cols_num(section), '1')

    def test_restore_body_column_section_supports_three_columns(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 3}):
            section = docx_pipe._restore_body_column_section(docx)
        self.assertEqual(self._cols_num(section), '3')

    def test_render_tables_skips_section_switch_when_body_is_single_column(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 1}), \
             patch('packtools.sps.formats.pdf.pipeline.docx.docx_renderer.table.add_table'):
            sections_before = len(docx.sections)
            docx_pipe._render_tables(docx, [{'layout': pdf_enum.SINGLE_COLUMN_PAGE_LABEL}])
        self.assertEqual(len(docx.sections), sections_before)

    def test_render_tables_switches_section_when_body_has_multiple_columns(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 2}), \
             patch('packtools.sps.formats.pdf.pipeline.docx.docx_renderer.table.add_table'):
            sections_before = len(docx.sections)
            docx_pipe._render_tables(docx, [{'layout': pdf_enum.SINGLE_COLUMN_PAGE_LABEL}])
        self.assertEqual(len(docx.sections), sections_before + 2)

    def test_render_figures_skips_section_switch_when_body_is_single_column(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 1}), \
             patch('packtools.sps.formats.pdf.pipeline.docx.docx_renderer.figure.add_figure'):
            sections_before = len(docx.sections)
            docx_pipe._render_figures(docx, [{'layout': pdf_enum.SINGLE_COLUMN_PAGE_LABEL}])
        self.assertEqual(len(docx.sections), sections_before)

    def test_render_figures_switches_section_when_body_has_multiple_columns(self):
        docx = Document()
        with patch.dict(pdf_enum.PAGE_ATTRIBUTES, {'default_column_count': 2}), \
             patch('packtools.sps.formats.pdf.pipeline.docx.docx_renderer.figure.add_figure'):
            sections_before = len(docx.sections)
            docx_pipe._render_figures(docx, [{'layout': pdf_enum.SINGLE_COLUMN_PAGE_LABEL}])
        self.assertEqual(len(docx.sections), sections_before + 2)


if __name__ == "__main__":
    unittest.main()
